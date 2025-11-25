import os
import re
import json
import tempfile
import requests
import subprocess
import logging
import sys
import webbrowser
import threading
import queue
import csv
import zipfile
from datetime import datetime
from pathlib import Path

import gradio as gr
import uvicorn
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

OLLAMA_MODEL = "gemma3:4b"

# -------------------------
# Hjälpfunktioner: GUI error dialogs
# -------------------------

def show_error_dialog(title, message):
    """
    Visa ett felmeddelande i en GUI-dialog (för windowed mode).
    Använder tkinter för att visa meddelandet.
    """
    try:
        import tkinter as tk
        from tkinter import messagebox
        
        # Skapa en dold root-window
        root = tk.Tk()
        root.withdraw()
        
        # Visa felmeddelandet
        messagebox.showerror(title, message)
        root.destroy()
    except Exception:
        # Om GUI inte fungerar, skriv till stderr
        print(f"ERROR: {title}\n{message}", file=sys.stderr)


def wait_for_enter():
    """
    Vänta på ENTER om programmet körs i en riktig terminal.
    I PyInstaller --noconsole / GUI-läge finns ofta ingen stdin.
    """
    try:
        if sys.stdin and sys.stdin.isatty():
            input("\nTryck ENTER för att stänga fönstret...")
    except Exception:
        # Ignorera alla fel här – vi vill inte krascha bara för att stdin saknas
        pass


# -------------------------
# Ollama-modellhantering
# -------------------------

def ensure_ollama_model(model=OLLAMA_MODEL):
    """
    Säkerställ att Ollama är igång, att modellen finns lokalt
    och värm upp modellen så att första riktiga anropet inte time:ar.
    """
    # 1. Kolla att Ollama svarar
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=10)
        resp.raise_for_status()
    except Exception as e:
        raise RuntimeError(
            "Ollama verkar inte vara igång.\n"
            "Starta Ollama-appen (eller Ollama-tjänsten) och försök igen."
        ) from e

    # 2. Kolla om modellen redan finns
    try:
        data = resp.json()
        models = data.get("models", [])
        names = [m.get("name") for m in models if isinstance(m, dict)]
    except Exception:
        names = []

    if model not in names:
        print(f"🧩 Modellen '{model}' saknas. Försöker ladda ner med 'ollama pull {model}'...")
        try:
            # Detta kräver att 'ollama' finns i PATH på datorn
            subprocess.run(
                ["ollama", "pull", model],
                check=True,
            )
            print(f"✅ Klar: modellen '{model}' nedladdad.")
        except Exception as e:
            raise RuntimeError(
                f"Kunde inte ladda ner modellen '{model}'.\n"
                f"Kontrollera att Ollama är korrekt installerat och att kommandot "
                f"'ollama pull {model}' fungerar.\n"
                f"Teknisk info: {e}"
            ) from e

    # 3. Warmup – litet generate-anrop så modellen laddas i RAM innan första riktiga jobbet
    print("🔥 Värmer upp modellen (första gången kan detta ta en stund)...")
    try:
        warmup_resp = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": "Värm upp modellen kort.",
                "stream": False,
                "options": {
                    "num_predict": 1  # bara en token, vi bryr oss inte om svaret
                },
            },
            timeout=900,
        )
        warmup_resp.raise_for_status()
        print("✅ Modellen är uppvärmd och redo.\n")
    except Exception as e:
        raise RuntimeError(
            "Kunde prata med Ollama men misslyckades med att värma upp modellen.\n"
            "Testa att köra 'ollama run gemma3:4b' manuellt en gång och prova sedan igen.\n"
            f"Teknisk info: {e}"
        ) from e


# -------------------------
# Batch Processing Queue
# -------------------------

class ProcessingJob:
    """Representerar ett enskilt filbearbetningsjobb."""
    
    def __init__(self, file_path, original_filename):
        self.file_path = file_path
        self.original_filename = original_filename
        self.status = "Köad"  # Köad, Bearbetar, Klar, Fel
        self.progress = 0
        self.error_message = ""
        self.entity_mapping = {}  # {original: pseudonym}
        self.output_path = None
        self.csv_path = None
        self.start_time = None
        self.end_time = None
    
    def to_dict(self):
        """Konvertera till dictionary för UI-visning."""
        return {
            "Filnamn": self.original_filename,
            "Status": self.status,
            "Progress": f"{self.progress}%",
            "Meddelande": self.error_message if self.error_message else "—"
        }


class BatchQueue:
    """Hanterar kö av filer för batchbearbetning."""
    
    def __init__(self, max_concurrent=2):
        self.jobs = []
        self.max_concurrent = max_concurrent
        self.lock = threading.Lock()
        self.processing_threads = []
        self.stop_flag = threading.Event()
    
    def add_job(self, file_path, original_filename):
        """Lägg till ett jobb i kön."""
        with self.lock:
            job = ProcessingJob(file_path, original_filename)
            self.jobs.append(job)
            return job
    
    def get_queued_jobs(self):
        """Hämta alla jobb som väntar på bearbetning."""
        with self.lock:
            return [j for j in self.jobs if j.status == "Köad"]
    
    def get_processing_jobs(self):
        """Hämta alla jobb som bearbetas just nu."""
        with self.lock:
            return [j for j in self.jobs if j.status == "Bearbetar"]
    
    def get_all_jobs_status(self):
        """Hämta status för alla jobb."""
        with self.lock:
            return [j.to_dict() for j in self.jobs]
    
    def get_completed_jobs(self):
        """Hämta alla färdiga jobb."""
        with self.lock:
            return [j for j in self.jobs if j.status == "Klar"]
    
    def start_processing(self):
        """Starta bearbetning av kön."""
        self.stop_flag.clear()
        # Starta worker-trådar
        for i in range(self.max_concurrent):
            thread = threading.Thread(target=self._worker, daemon=True)
            thread.start()
            self.processing_threads.append(thread)
    
    def stop_processing(self):
        """Stoppa bearbetning av kön."""
        self.stop_flag.set()
    
    def _worker(self):
        """Worker-tråd som bearbetar jobb från kön."""
        while not self.stop_flag.is_set():
            # Hämta nästa jobb
            job = None
            with self.lock:
                queued = [j for j in self.jobs if j.status == "Köad"]
                if queued:
                    job = queued[0]
                    job.status = "Bearbetar"
                    job.start_time = datetime.now()
            
            if job is None:
                # Inget jobb att bearbeta, vänta lite
                threading.Event().wait(1)
                continue
            
            # Bearbeta jobbet
            try:
                self._process_job(job)
                job.status = "Klar"
                job.progress = 100
                print(f"✅ Klar: {job.original_filename}")
            except Exception as e:
                job.status = "Fel"
                job.error_message = str(e)
                job.progress = 0
                print(f"❌ Fel vid bearbetning av {job.original_filename}: {e}")
                import traceback
                traceback.print_exc()
            finally:
                job.end_time = datetime.now()
    
    def _process_job(self, job):
        """Bearbeta ett enskilt jobb."""
        # Läs dokument
        job.progress = 10
        
        # Skapa en mock file object
        class FileObj:
            def __init__(self, path):
                self.name = path
        
        file_obj = FileObj(job.file_path)
        content, ext = read_document(file_obj)
        
        # Extrahera text för analys
        job.progress = 20
        text = extract_text_from_content(content, ext)
        
        # Analysera med Ollama
        job.progress = 30
        ai_result = call_ollama_extract(text)
        persons = ai_result.get("persons", [])
        pnr_list = ai_result.get("personnummer", []) or []
        phone_list = ai_result.get("phones", []) or []
        
        # Komplettera med regex
        job.progress = 50
        rx_entities = regex_find_entities(text)
        for e in rx_entities:
            if e["type"] == "personnummer":
                pnr_list.append(e["value"])
            elif e["type"] == "phone":
                phone_list.append(e["value"])
        
        # Bygg mapping
        job.progress = 60
        mapping = {}
        counters = {"name": 1, "personnummer": 1, "phone": 1}
        
        # Personer
        for person in persons:
            pseudo = f"Person_{counters['name']:03d}"
            counters["name"] += 1
            for mention in person.get("mentions", []):
                if mention:
                    mapping[mention] = pseudo
        
        # Personnummer
        for pnr in set(pnr_list):
            mapping[pnr] = f"PNR_{counters['personnummer']:03d}"
            counters["personnummer"] += 1
        
        # Telefoner
        for phone in set(phone_list):
            mapping[phone] = f"TEL_{counters['phone']:03d}"
            counters["phone"] += 1
        
        job.entity_mapping = mapping
        
        # Pseudonymisera
        job.progress = 80
        new_content = pseudonymize_content(content, ext, mapping)
        
        # Skriv output
        job.progress = 90
        output_path = write_output_file(new_content, ext)
        job.output_path = output_path
        
        # Skapa CSV med mappingen
        csv_path = create_mapping_csv(mapping, job.original_filename)
        job.csv_path = csv_path
        
        job.progress = 100


def create_mapping_csv(mapping, original_filename):
    """Skapa CSV-fil med entity mappings."""
    fd, csv_path = tempfile.mkstemp(suffix=".csv")
    
    with os.fdopen(fd, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Typ", "Original", "Pseudonym", "Filnamn"])
        
        for original, pseudo in mapping.items():
            # Gissa typ baserat på pseudonym
            if pseudo.startswith("Person_"):
                typ = "name"
            elif pseudo.startswith("PNR_"):
                typ = "personnummer"
            elif pseudo.startswith("TEL_"):
                typ = "phone"
            else:
                typ = "unknown"
            
            writer.writerow([typ, original, pseudo, original_filename])
    
    return csv_path


def create_download_package(jobs):
    """Skapa zip-fil med alla resultat."""
    fd, zip_path = tempfile.mkstemp(suffix=".zip")
    os.close(fd)
    
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for job in jobs:
            if job.status == "Klar" and job.output_path:
                # Lägg till pseudonymiserad fil
                base_name = Path(job.original_filename).stem
                ext = Path(job.original_filename).suffix
                zipf.write(job.output_path, f"{base_name}_pseudonymiserad{ext}")
                
                # Lägg till CSV
                if job.csv_path:
                    zipf.write(job.csv_path, f"{base_name}_mapping.csv")
    
    return zip_path


# Global batch queue instance
batch_queue = BatchQueue(max_concurrent=2)


# -------------------------
# Hjälpfunktioner för text
# -------------------------

def read_pdf_structured(filename):
    """
    Läs PDF och bevara struktur (sidor och stycken).
    Returnerar lista av stycken med metadata.
    """
    paragraphs = []
    try:
        with open(filename, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    # Dela upp i stycken baserat på dubbla radbrytningar
                    page_paragraphs = text.split('\n\n')
                    for para_text in page_paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            paragraphs.append({
                                'text': para_text,
                                'style': 'normal',
                                'page': page_num + 1
                            })
    except Exception as e:
        raise ValueError(f"Fel vid läsning av PDF: {e}")
    
    return paragraphs


def read_docx_structured(filename):
    """
    Läs DOCX och bevara struktur (stycken med formatering).
    Returnerar lista av stycken med metadata.
    """
    doc = Document(filename)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Identifiera stil
        style = 'normal'
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            style = para.style.name.lower()  # 'heading 1', 'heading 2', etc.
        
        # Samla formateringsinformation från runs
        runs_info = []
        for run in para.runs:
            runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline
            })
        
        paragraphs.append({
            'text': text,
            'style': style,
            'runs': runs_info,
            'alignment': para.alignment
        })
    
    return paragraphs


def read_document(file_obj):
    """
    Läs text ur .txt, .docx eller .pdf.
    Returnerar (content, extension) där:
    - content är sträng för .txt
    - content är lista av stycken med metadata för .docx och .pdf
    """
    filename = file_obj.name
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        with open(filename, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), ".txt"
    elif ext == ".docx":
        paragraphs = read_docx_structured(filename)
        return paragraphs, ".docx"
    elif ext == ".pdf":
        paragraphs = read_pdf_structured(filename)
        return paragraphs, ".pdf"
    else:
        raise ValueError("Endast .txt, .docx och .pdf stöds.")


def merge_persons(persons):
    """
    Slår ihop personposter som sannolikt är samma person,
    baserat på canonical_name och överlappande mentions.
    """
    merged = []
    for p in persons:
        canonical = (p.get("canonical_name") or "").strip()
        mentions = [
            m.strip()
            for m in (p.get("mentions") or [])
            if isinstance(m, str) and m.strip()
        ]

        if not canonical and not mentions:
            continue

        placed = False
        for m in merged:
            m_canonical = m.get("canonical_name") or ""
            m_mentions = m.get("mentions") or []

            # Samma canonical_name
            if canonical and m_canonical and canonical == m_canonical:
                all_mentions = set(m_mentions) | set(mentions)
                if canonical:
                    all_mentions.add(canonical)
                m["mentions"] = list(all_mentions)
                placed = True
                break

            # Överlappande mentions
            if set(mentions) & set(m_mentions):
                all_mentions = set(m_mentions) | set(mentions)
                if not m_canonical and canonical:
                    m["canonical_name"] = canonical
                if canonical:
                    all_mentions.add(canonical)
                m["mentions"] = list(all_mentions)
                placed = True
                break

        if not placed:
            merged.append(
                {
                    "canonical_name": canonical,
                    "mentions": mentions if mentions else ([canonical] if canonical else []),
                }
            )

    return merged


def extract_text_from_content(content, ext):
    """
    Extrahera ren text från innehåll (sträng eller strukturerad data).
    Används för AI-analys.
    """
    if ext == ".txt":
        return content
    elif ext in [".docx", ".pdf"]:
        # content är lista av stycken
        if isinstance(content, list):
            return "\n\n".join(p.get('text', '') for p in content)
        return str(content)
    return str(content)


def call_ollama_extract(text, model=OLLAMA_MODEL):
    """
    Anropar Ollama för att:
    - hitta personer och alla sätt de omnämns på
    - hitta personnummer/födelsedatum
    - hitta telefonnummer

    Returnerar ett dict:
    {
      "persons": [...],
      "personnummer": [...],
      "phones": [...]
    }
    """
    # Begränsa längden för säkerhets skull (kan justeras upp vid behov)
    limited_text = text[:8000]

    prompt = f"""
Du analyserar svensk löptext och extraherar personuppgifter (PII).

Din uppgift är att hitta:
- alla personer och alla sätt de omnämns på
- alla personnummer/födelsedatum kopplade till personer
- alla telefonnummer

Viktiga regler:

1. Personer
- Identifiera alla personer i texten, även om de bara nämns med förnamn.
- "canonical_name" ska vara det mest kompletta namnet som finns i texten för personen
  (t.ex. "Emma Sandberg" om både "Emma" och "Emma Sandberg" förekommer,
   annars bara "Emma" om förnamnet är det enda som finns).
- Om "canonical_name" saknas för en person men det finns "mentions",
  ska du låta den längsta mention-strängen representera personen.
- "mentions" ska innehålla alla exakta textsträngar som syftar på personen,
  t.ex. ["Emma Sandberg", "Emma"]. Använd exakt stavning från texten.
- Tänk på att namn kan ha ett genitiv-s "Emmas", det ska ändå ingå i "mentions" för den personen.
- Ta inte med pronomen (han, hon, de, honom, henne, etc.).
- Ta med smeknamn och namn i citat, t.ex. "Oskar" i "den där killen som jobbar med API-nycklarna, Oskar".
- Om texten innehåller en beskrivande fras som tydligt syftar på en specifik person,
  t.ex. "den där killen som jobbar med API-nycklarna" eller "chefen som alltid ringer vid sju",
  ska hela frasen också ingå i "mentions" för den personen.
- Om en fras både innehåller en beskrivning och ett namn, t.ex.
  "den där killen som jobbar med API-nycklarna, Oskar, tror jag?",
  ska både hela frasen och namnet ("Oskar") finnas med i "mentions".
- Hitta inte på egna namn som inte finns i texten.

2. Personnummer och födelsedatum
- Lista alla strängar som ser ut som svenska personnummer eller födelsedatum kopplade till en person.
- Ta med format som:
  - YYYYMMDD-XXXX
  - YYYYMMDDXXXX
  - YYMMDD-XXXX
  - YYMMDDXXXX
  - YYYY-MM-DD (om det är tydligt kopplat till en person, t.ex. står efter ett namn eller inom parentes).
- Returnera strängarna exakt som de står i texten.

3. Telefonnummer
- Lista alla strängar som ser ut som telefonnummer, t.ex.:
  - 07XXXXXXXX
  - 0X-XXX XX XX
  - 0XX-XXXXXXX
  - +46 7X XXX XX XX
  - +467XXXXXXXX
- Inkludera varianter med mellanslag och bindestreck.
- Exkludera:
  - klockslag (t.ex. "07.43", "16.28")
  - fakturanummer
  - versionsnummer (t.ex. "3.7.14")
  - andra rena nummer som inte tydligt är telefonnummer.

4. Allmänt
- Alla värden i "mentions", "personnummer" och "phones" ska vara exakta substrängar från texten.
- Hitta så många korrekta träffar som möjligt. Missa hellre inget än att vara överförsiktig.
- Svara med välformad JSON som går att parsa direkt.

Text:
\"\"\"{limited_text}\"\"\"


Svara ENBART med giltig JSON på formen:

{{
  "persons": [
    {{
      "canonical_name": "Anna Björk",
      "mentions": ["Anna Björk", "Anna"]
    }}
  ],
  "personnummer": ["1986-05-19", "YYYYMMDD-XXXX"],
  "phones": ["+46 70 123 45 67", "070-123 45 67"]
}}

Inga kommentarer, ingen extra text, endast JSON.
"""

    resp = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": model,
            "prompt": prompt,
            "stream": False,
        },
        timeout=900,
    )
    resp.raise_for_status()
    data = resp.json()
    raw = data.get("response", "").strip()

    # Enkel städning om modellen skulle råka lägga till ```json ``` runt
    if raw.startswith("```"):
        raw = raw.strip("`")
        # ta bort ev. "json" i första raden
        if raw.lower().startswith("json"):
            raw = raw[4:].strip()
        # hitta första { och sista }
        start = raw.find("{")
        end = raw.rfind("}")
        if start != -1 and end != -1:
            raw = raw[start:end+1].strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        # Om modellen spårar ur, returnera tom struktur
        return {"persons": [], "personnummer": [], "phones": []}

    persons = parsed.get("persons", []) or []
    personnummer = parsed.get("personnummer", []) or []
    phones = parsed.get("phones", []) or []

    # Städa upp personer
    clean_persons = []
    for p in persons:
        canonical = (p.get("canonical_name") or "").strip()
        mentions = p.get("mentions") or []
        mentions = [m.strip() for m in mentions if isinstance(m, str) and m.strip()]

        # Om canonical saknas men vi har mentions → välj längsta mention som canonical
        if not canonical and mentions:
            canonical = max(mentions, key=len)

        if not canonical and not mentions:
            continue

        if canonical and canonical not in mentions:
            mentions.insert(0, canonical)

        clean_persons.append(
            {
                "canonical_name": canonical,
                "mentions": mentions,
            }
        )

    # Slå ihop dubbletter av samma person (t.ex. "Oskar" + "Oskar Nyblom")
    clean_persons = merge_persons(clean_persons)

    personnummer = [
        s.strip() for s in personnummer if isinstance(s, str) and s.strip()
    ]
    phones = [s.strip() for s in phones if isinstance(s, str) and s.strip()]

    return {
        "persons": clean_persons,
        "personnummer": personnummer,
        "phones": phones,
    }


def regex_find_entities(text):
    """Komplettera med regex för personnummer/telefonnummer."""
    entities = []

    # Svenska personnummer: 6+4 eller 8+4 samt födelsedatum YYYY-MM-DD
    pnr_pattern = r"\b\d{6}[-+]\d{4}\b|\b\d{8}-\d{4}\b|\b\d{4}-\d{2}-\d{2}\b"
    for match in re.findall(pnr_pattern, text):
        entities.append({"type": "personnummer", "value": match})

    # Telefonnr: grovt mönster för +46 / 0 följt av siffror och mellanslag
    phone_pattern = r"\b(?:\+46|0)\d[\d\s\-]{6,}\d\b"
    for match in re.findall(phone_pattern, text):
        entities.append({"type": "phone", "value": match})

    return entities


def deduplicate_entities(entities):
    """Ta bort dubbletter (type + value)."""
    seen = set()
    unique = []
    for e in entities:
        key = (e["type"], e["value"])
        if key not in seen:
            seen.add(key)
            unique.append(e)
    return unique


# -------------------------
# Pseudonymisering
# -------------------------

def build_mapping_from_table(table_rows):
    """
    Tar in tabellen från Gradio (kan vara en pandas DataFrame eller en lista av listor)
    och bygger en mapping {original: pseudonym}.

    - Kolumner förväntas vara: [Typ, Original, Pseudonym]
    - Om Pseudonym är tom genereras den automatiskt.
    """

    # 1. Konvertera ev. DataFrame -> lista av listor
    if hasattr(table_rows, "values"):  # pandas DataFrame
        rows = table_rows.values.tolist()
    else:
        rows = table_rows or []

    mapping = {}
    counters = {"name": 1, "personnummer": 1, "phone": 1}

    if rows is None or len(rows) == 0:
        return mapping

    for row in rows:
        if not row:
            continue

        # Se till att vi har minst 2 kolumner (typ + original)
        # Pseudonymkolumnen kan saknas, då blir den tom.
        t = row[0] if len(row) > 0 else ""
        original = row[1] if len(row) > 1 else ""
        pseudo = row[2] if len(row) > 2 else ""

        # Om något av dessa är NaN/None etc, gör dem till tomma strängar
        t = "" if t is None else str(t).strip()
        original = "" if original is None else str(original).strip()
        pseudo = "" if pseudo is None else str(pseudo).strip()

        if not original:
            continue  # hoppa rader utan originalvärde

        if not t:
            t = "name"  # default

        # Auto-generera pseudonym om tom
        if not pseudo:
            if t == "name":
                pseudo = f"Person_{counters['name']:03d}"
                counters["name"] += 1
            elif t == "personnummer":
                pseudo = f"PNR_{counters['personnummer']:03d}"
                counters["personnummer"] += 1
            elif t == "phone":
                pseudo = f"TEL_{counters['phone']:03d}"
                counters["phone"] += 1
            else:
                pseudo = f"X_{counters['name']:03d}"
                counters["name"] += 1

        mapping[original] = pseudo

    return mapping


def pseudonymize_text(text, mapping):
    """
    Ersätt alla förekomster av original med pseudonym i ren text.
    Lite försiktigare för namn (ordgränser).
    """
    new_text = text

    # Sortera så att längre strängar ersätts först (minskar risk för delmatchning)
    items = sorted(mapping.items(), key=lambda kv: len(kv[0]), reverse=True)

    for original, pseudo in items:
        # Heuristik: om det ser ut som ett personnummer eller telefon, byt rakt av
        if re.match(r"^\d{6}[-+]\d{4}$|^\d{8}-\d{4}$|^\d{4}-\d{2}-\d{2}$", original):
            pattern = re.escape(original)
        elif re.match(r"^(?:\+46|0)\d", original.replace(" ", "").replace("-", "")):
            pattern = re.escape(original)
        else:
            # Behandla som namn/fras – omge med ordgränser
            pattern = r"\b" + re.escape(original) + r"\b"

        new_text = re.sub(pattern, pseudo, new_text)

    return new_text


def pseudonymize_structured(content, mapping):
    """
    Pseudonymisera strukturerat innehåll (lista av stycken).
    Bevarar formatering och metadata.
    """
    new_content = []
    
    for para in content:
        new_para = para.copy()
        
        # Pseudonymisera texten
        new_para['text'] = pseudonymize_text(para['text'], mapping)
        
        # Om det finns runs (DOCX), pseudonymisera varje run
        if 'runs' in para:
            new_runs = []
            for run in para['runs']:
                new_run = run.copy()
                new_run['text'] = pseudonymize_text(run['text'], mapping)
                new_runs.append(new_run)
            new_para['runs'] = new_runs
        
        new_content.append(new_para)
    
    return new_content


def pseudonymize_content(content, ext, mapping):
    """
    Pseudonymisera innehåll baserat på filtyp.
    Returnerar pseudonymiserat innehåll i samma format som input.
    """
    if ext == ".txt":
        return pseudonymize_text(content, mapping)
    elif ext in [".docx", ".pdf"]:
        return pseudonymize_structured(content, mapping)
    else:
        return pseudonymize_text(str(content), mapping)


def write_docx_document(content, filename):
    """
    Skriv strukturerat innehåll till DOCX-fil med bevarad formatering.
    """
    doc = Document()
    
    for para_data in content:
        text = para_data.get('text', '')
        style = para_data.get('style', 'normal')
        runs = para_data.get('runs', [])
        alignment = para_data.get('alignment')
        
        # Skapa stycke med rätt stil
        if style.startswith('heading'):
            # Extrahera heading-nivå
            try:
                level = style.split()[-1]  # 'heading 1' -> '1'
                para = doc.add_heading('', level=int(level))  # Skapa tom heading
            except:
                para = doc.add_paragraph()
        else:
            para = doc.add_paragraph()
        
        # Lägg alltid till text via runs (för både headings och normala stycken)
        if runs:
            for run_data in runs:
                run = para.add_run(run_data.get('text', ''))
                if run_data.get('bold'):
                    run.bold = True
                if run_data.get('italic'):
                    run.italic = True
                if run_data.get('underline'):
                    run.underline = True
        else:
            # Fallback: lägg till texten som en enda run
            para.add_run(text)
        
        # Sätt justering om den finns
        if alignment is not None:
            para.alignment = alignment
    
    doc.save(filename)


def write_pdf_document(content, filename):
    """
    Skriv strukturerat innehåll till PDF-fil med grundläggande formatering.
    """
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Skapa anpassade stilar
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12,
        textColor='black'
    )
    
    normal_style = styles['Normal']
    
    story = []
    
    for para_data in content:
        text = para_data.get('text', '')
        style = para_data.get('style', 'normal')
        
        # Escapea special characters för reportlab
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Välj stil baserat på paragraftyp
        if style.startswith('heading'):
            para = Paragraph(text, heading_style)
        else:
            para = Paragraph(text, normal_style)
        
        story.append(para)
        story.append(Spacer(1, 0.2 * inch))
    
    doc.build(story)


def write_output_file(content, ext):
    """
    Skriv innehåll till fil baserat på filtyp.
    Returnerar sökväg till temporär fil.
    """
    if ext == ".txt":
        # Ren text
        fd, path = tempfile.mkstemp(suffix=".txt")
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(content)
        return path
    
    elif ext == ".docx":
        # Strukturerad DOCX
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)  # Stäng fd, Document() öppnar filen själv
        write_docx_document(content, path)
        return path
    
    elif ext == ".pdf":
        # Strukturerad PDF
        fd, path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        write_pdf_document(content, path)
        return path
    
    else:
        # Fallback till text
        fd, path = tempfile.mkstemp(suffix=".txt")
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(str(content))
        return path


# -------------------------
# Gradio-callbacks
# -------------------------

def analyze_document(file_obj):
    if file_obj is None:
        return [], "Ingen fil uppladdad ännu."

    try:
        content, ext = read_document(file_obj)
    except Exception as e:
        return [], f"Fel vid läsning av dokument: {e}"

    # Extrahera ren text för AI-analys
    text = extract_text_from_content(content, ext)

    # Hämta strukturerad info från Ollama
    ai_result = call_ollama_extract(text)
    persons = ai_result.get("persons", [])
    pnr_list = ai_result.get("personnummer", []) or []
    phone_list = ai_result.get("phones", []) or []

    # Komplettera personnummer/telefoner med regex (säkerhetsbälte)
    rx_entities = regex_find_entities(text)
    for e in rx_entities:
        if e["type"] == "personnummer":
            pnr_list.append(e["value"])
        elif e["type"] == "phone":
            phone_list.append(e["value"])

    rows = []
    status_parts = []

    # 1. Personer: en pseudonym per person, samma för alla omnämnanden
    person_counter = 1
    total_mentions = 0
    for person in persons:
        pseudo = f"Person_{person_counter:03d}"
        person_counter += 1

        mentions = person.get("mentions", [])
        added = set()
        for mention in mentions:
            if not mention:
                continue
            if mention in added:
                continue
            added.add(mention)
            rows.append(["name", mention, pseudo])
            total_mentions += 1

    if persons:
        status_parts.append(
            f"hittade {len(persons)} person(er) med totalt {total_mentions} namn-omnämnanden"
        )

    # 2. Personnummer – en rad per unikt nummer (pseudonym tom -> auto-genereras senare)
    unique_pnr = sorted(set(pnr_list))
    for p in unique_pnr:
        rows.append(["personnummer", p, ""])
    if unique_pnr:
        status_parts.append(f"{len(unique_pnr)} personnummer")

    # 3. Telefoner – en rad per unikt nummer
    unique_phones = sorted(set(phone_list))
    for t in unique_phones:
        rows.append(["phone", t, ""])
    if unique_phones:
        status_parts.append(f"{len(unique_phones)} telefonnummer")

    if not rows:
        msg = (
            "Inga namn, personnummer eller telefonnummer hittades automatiskt. "
            "Du kan ändå lägga till rader manuellt."
        )
    else:
        msg = (
            "Träffar: "
            + ", ".join(status_parts)
            + ". Justera listan vid behov och klicka sedan på 'Skapa pseudonymiserat dokument'."
        )

    return rows, msg


def run_pseudonymization(file_obj, table_rows):
    if file_obj is None:
        return "Ingen fil.", None

    try:
        content, ext = read_document(file_obj)
    except Exception as e:
        return f"Fel vid läsning av dokument: {e}", None

    # Bygg mapping från tabellen (oavsett om det är lista eller DataFrame)
    mapping = build_mapping_from_table(table_rows)

    if not mapping:
        return (
            "Ingen mapping kunde skapas från tabellen. "
            "Kontrollera att du klickat på 'Analysera dokument' först och att tabellen inte är helt tom.",
            None,
        )

    # Pseudonymisera innehållet (bevarar struktur för DOCX/PDF)
    new_content = pseudonymize_content(content, ext, mapping)
    
    # Skriv till fil i samma format som input
    out_path = write_output_file(new_content, ext)
    
    # Skapa förhandsgranskning
    if ext == ".txt":
        preview = new_content[:1000]
    else:
        # För strukturerade dokument, visa text från första styckena
        preview_text = extract_text_from_content(new_content, ext)
        preview = preview_text[:1000]

    return preview, out_path


# -------------------------
# Batch Processing Callbacks
# -------------------------

def upload_batch_files(files):
    """Ladda upp flera filer till batch-kön."""
    if not files:
        return "Inga filer uppladdade.", []
    
    # Rensa gamla jobb om kön är tom
    if not batch_queue.get_queued_jobs() and not batch_queue.get_processing_jobs():
        batch_queue.jobs.clear()
    
    # Lägg till filer i kön
    for file in files:
        batch_queue.add_job(file.name, os.path.basename(file.name))
    
    status = f"✅ {len(files)} fil(er) tillagda i kön."
    return status, batch_queue.get_all_jobs_status()


def start_batch_processing():
    """Starta batchbearbetning."""
    queued = batch_queue.get_queued_jobs()
    if not queued:
        return "Inga filer i kön att bearbeta.", batch_queue.get_all_jobs_status()
    
    batch_queue.start_processing()
    return f"🚀 Bearbetning startad! {len(queued)} fil(er) i kön.", batch_queue.get_all_jobs_status()


def stop_batch_processing():
    """Stoppa batchbearbetning."""
    batch_queue.stop_processing()
    return "⏸️ Bearbetning stoppad.", batch_queue.get_all_jobs_status()


def refresh_batch_status():
    """Uppdatera status för batch-jobb."""
    jobs_status = batch_queue.get_all_jobs_status()
    
    completed = len([j for j in batch_queue.jobs if j.status == "Klar"])
    total = len(batch_queue.jobs)
    processing = len(batch_queue.get_processing_jobs())
    queued = len(batch_queue.get_queued_jobs())
    
    summary = f"📊 Status: {completed}/{total} klara | {processing} bearbetas | {queued} i kö"
    
    return summary, jobs_status


def download_batch_results():
    """Skapa och returnera zip-fil med alla resultat."""
    completed_jobs = batch_queue.get_completed_jobs()
    
    if not completed_jobs:
        return None
    
    zip_path = create_download_package(completed_jobs)
    return zip_path


def clear_batch_queue():
    """Rensa batch-kön."""
    batch_queue.stop_processing()
    batch_queue.jobs.clear()
    return "🗑️ Kön rensad.", []


# -------------------------
# Bygg Gradio-gränssnitt
# -------------------------

with gr.Blocks(title="Lokal Pseudonymiserare") as demo:
    gr.Markdown(
        "# Lokal pseudonymiserare\n"
        "Pseudonymisera dokument med lokal AI (Ollama). Välj mellan **Enkel fil** eller **Batch-läge** för flera filer."
    )
    
    with gr.Tabs():
        # Tab 1: Single File Mode (original)
        with gr.Tab("📄 Enkel fil"):
            gr.Markdown(
                "Ladda upp ett dokument, granska och redigera identifierade entiteter, "
                "och generera ett pseudonymiserat dokument. **Dokumentstrukturen bevaras** för DOCX och PDF."
            )
            
            with gr.Row():
                file_in = gr.File(
                    label="Ladda upp dokument (.txt, .docx eller .pdf)",
                    file_types=[".txt", ".docx", ".pdf"],
                )
                analyze_btn = gr.Button("Analysera dokument")

            entities_table = gr.Dataframe(
                headers=["Typ", "Original", "Pseudonym"],
                datatype=["str", "str", "str"],
                row_count=(0, "dynamic"),
                interactive=True,
                label="Identifierade (och redigerbara) entiteter",
            )

            status_box = gr.Markdown("Statusmeddelanden visas här.")

            analyze_btn.click(
                fn=analyze_document,
                inputs=file_in,
                outputs=[entities_table, status_box],
            )

            gr.Markdown("### Steg 2: Pseudonymisera dokumentet")
            pseud_btn = gr.Button("Skapa pseudonymiserat dokument")

            preview_box = gr.Textbox(
                label="Förhandsgranskning (första 1000 tecken)",
                lines=15,
            )
            download_out = gr.File(label="Ladda ner pseudonymiserad fil (samma format som input)")

            pseud_btn.click(
                fn=run_pseudonymization,
                inputs=[file_in, entities_table],
                outputs=[preview_box, download_out],
            )
        
        # Tab 2: Batch Mode
        with gr.Tab("📦 Batch-läge"):
            gr.Markdown(
                "### Batch-bearbetning av flera filer\n"
                "Ladda upp flera filer samtidigt (upp till 50+). Filerna bearbetas automatiskt i kö "
                "(max 2 samtidigt för att inte överbelasta datorn). **OBS:** Entiteter identifieras "
                "automatiskt utan manuell granskning."
            )
            
            with gr.Row():
                batch_files_in = gr.File(
                    label="Ladda upp flera dokument (.txt, .docx eller .pdf)",
                    file_types=[".txt", ".docx", ".pdf"],
                    file_count="multiple"
                )
            
            with gr.Row():
                upload_batch_btn = gr.Button("📥 Lägg till i kö", variant="primary")
                start_batch_btn = gr.Button("▶️ Starta bearbetning", variant="primary")
                stop_batch_btn = gr.Button("⏸️ Stoppa")
                clear_batch_btn = gr.Button("🗑️ Rensa kö")
            
            batch_status_text = gr.Markdown("Ingen batch aktiv.")
            
            batch_status_table = gr.Dataframe(
                headers=["Filnamn", "Status", "Progress", "Meddelande"],
                datatype=["str", "str", "str", "str"],
                label="Batch-status",
                interactive=False,
            )
            
            with gr.Row():
                refresh_btn = gr.Button("🔄 Uppdatera status")
                download_batch_btn = gr.Button("💾 Ladda ner alla resultat (ZIP)", variant="primary")
            
            batch_download_out = gr.File(label="Batch-resultat (ZIP)")
            
            # Batch callbacks
            upload_batch_btn.click(
                fn=upload_batch_files,
                inputs=batch_files_in,
                outputs=[batch_status_text, batch_status_table],
            )
            
            start_batch_btn.click(
                fn=start_batch_processing,
                outputs=[batch_status_text, batch_status_table],
            )
            
            stop_batch_btn.click(
                fn=stop_batch_processing,
                outputs=[batch_status_text, batch_status_table],
            )
            
            clear_batch_btn.click(
                fn=clear_batch_queue,
                outputs=[batch_status_text, batch_status_table],
            )
            
            refresh_btn.click(
                fn=refresh_batch_status,
                outputs=[batch_status_text, batch_status_table],
            )
            
            download_batch_btn.click(
                fn=download_batch_results,
                outputs=batch_download_out,
            )


# -------------------------
# Uvicorn-loggpatch för PyInstaller-miljö
# -------------------------

def patch_uvicorn_logging():
    """
    PyInstaller + uvicorn kan krascha när uvicorns DefaultFormatter försöker
    använda extra kwargs (t.ex. use_colors) mot logging.Formatter.
    Vi ersätter formatterna 'default' och 'access' med logging.Formatter och rensar bort
    alla nycklar som inte stöds.
    """
    try:
        from uvicorn.config import LOGGING_CONFIG
        
        # Endast tillåtna nycklar för logging.Formatter via dictConfig
        allowed_keys = {"()", "fmt", "datefmt", "style", "validate"}
        
        # Patcha både 'default' och 'access' formatters
        for formatter_name in ["default", "access"]:
            formatter = LOGGING_CONFIG.get("formatters", {}).get(formatter_name)
            if isinstance(formatter, dict):
                # Använd standard-Formatter istället för Uvicorns egen
                formatter["()"] = "logging.Formatter"
                
                # Ta bort nycklar som inte stöds
                for key in list(formatter.keys()):
                    if key not in allowed_keys:
                        formatter.pop(key, None)
    except Exception as e:
        # Om patchen misslyckas vill vi inte krascha programmet
        print(f"Kunde inte patcha uvicorn-logging: {e}")


if __name__ == "__main__":
    patch_uvicorn_logging()
    
    # Batch processing logs kommer fortfarande att skrivas (kan ses i log-fil)
    print("🚀 Startar Lokal Pseudonymiserare...")
    print("Öppna webbläsaren och gå till http://127.0.0.1:7860 om den inte öppnas automatiskt.")

    try:
        ensure_ollama_model()
    except RuntimeError as e:
        # Visa felmeddelande i GUI-dialog
        show_error_dialog(
            "Ollama-fel",
            str(e)
        )
        sys.exit(1)
    else:
        try:
            # Gradio startar servern och öppnar webbläsaren
            demo.launch(
                server_name="127.0.0.1",
                server_port=7860,
                inbrowser=True,   # öppna webbläsare automatiskt
                share=False,
            )
        except Exception as e:
            import traceback
            error_msg = f"Ett oväntat fel uppstod:\n\n{str(e)}\n\nSe konsolen för mer detaljer."
            show_error_dialog("Programfel", error_msg)
            traceback.print_exc()
            sys.exit(1)
